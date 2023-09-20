import json
from django.http import HttpResponse, JsonResponse
import docx2txt
from docx import Document
from docx.shared import Pt
import re
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import pdfkit



def extract_hyperlinks(run):
    hyperlinks = []
    for fld in run._element.findall('.//w:hyperlink', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
        href = fld.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}href")
        text = fld.text
        if href and text:
            hyperlinks.append((text, href))
    return hyperlinks

def extract_email_hyperlinks(run):
    text = run.text
    email_links = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
    return email_links

def apply_format(run, style):
    if style == 'BOLD':
        run.bold = True
    elif style == 'ITALIC':
        run.italic = True
    elif style == 'UNDERLINE':
        run.underline = True

def generate_docx_from_html_content(request):
    data = json.loads(request.body.decode('utf-8'))
    html_content = data.get("content")
    doc = Document()

    for block in html_content['blocks']:
        text = block['text']
        paragraph = doc.add_paragraph()

        if 'inlineStyleRanges' in block and block['inlineStyleRanges']:
            for inline_style in block['inlineStyleRanges']:
                offset = inline_style['offset']
                length = inline_style['length']
                style = inline_style['style']
                run = paragraph.add_run(text[offset:offset + length])
                apply_format(run, style)

            remaining_text = text[length:]
            if remaining_text:
                run = paragraph.add_run(remaining_text)
        else:
            run = paragraph.add_run(text)

    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save('EditedWordFile.docx')

    with open('EditedWordFile.docx', 'rb') as docx_file:
        response = HttpResponse(docx_file.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        response['Content-Disposition'] = 'attachment; filename=EditedWordFile.docx'
        return response



def generate_pdf_from_html_content(request):
    data = json.loads(request.body.decode('utf-8'))
    html_content = data.get("content")
    doc = Document()

    for block in html_content['blocks']:
        text = block['text']
        paragraph = doc.add_paragraph()

        if 'inlineStyleRanges' in block and block['inlineStyleRanges']:
            for inline_style in block['inlineStyleRanges']:
                offset = inline_style['offset']
                length = inline_style['length']
                style = inline_style['style']
                run = paragraph.add_run(text[offset:offset + length])
                apply_format(run, style)

            remaining_text = text[length:]
            if remaining_text:
                run = paragraph.add_run(remaining_text)
        else:
            run = paragraph.add_run(text)

    for paragraph in doc.paragraphs:
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save('EditedPdfFile.pdf')

    with open('EditedPdfFile.pdf', 'rb') as pdf_file:
        response = HttpResponse(pdf_file.read(), content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename=EditedPdfFile.pdf'
        return response

def parse_docx_to_html(request):
    if request.method == 'POST':
        uploaded_file = request.FILES['file']
        doc = Document(uploaded_file)
        html_output = []

        for paragraph in doc.paragraphs:
            html_paragraph = "<p>"
            for run in paragraph.runs:
                font_size = run.font.size
                if font_size:
                    html_paragraph += f'<span style="font-size:{font_size.pt}px;">'

                if run.bold:
                    html_paragraph += "<strong>"
                if run.italic:
                    html_paragraph += "<em>"

                if run.underline:
                    html_paragraph += "<u>"
                
                email_links = extract_email_hyperlinks(run)
                
                if email_links:
                    for email in email_links:
                        email_html = f'<a href="mailto:{email}">{email}</a>'
                        text = run.text.replace(email, email_html)
                        run.text = text

                html_paragraph += run.text

                if run.underline:
                    html_paragraph += "</u>"
                if run.italic:
                    html_paragraph += "</em>"
                if run.bold:
                    html_paragraph += "</strong>"
                if font_size:
                    html_paragraph += "</span>"

            html_paragraph += "</p>"
            html_output.append(html_paragraph)

        text = "\n".join(html_output)
        html_content = '<div>' + text.replace('\n', '<br>') + '</div>'
        return JsonResponse({'html_content': html_content})
    return JsonResponse({'error': 'Invalid request method'}, status=400)
